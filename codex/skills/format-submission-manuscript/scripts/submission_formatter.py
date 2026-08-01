#!/usr/bin/env python3
"""Deterministic publication-formatting layer for semantic manuscripts.

The module deliberately keeps research content immutable.  It reads a Markdown
manuscript plus a declared submission contract, writes a new run directory, and
records enough hashes and QA evidence to reproduce the conversion.
"""

from __future__ import annotations

import argparse
import contextlib
import copy
import datetime as dt
import hashlib
import json
import os
import platform
import re
import shutil
import subprocess
import sys
import tempfile
import textwrap
import uuid
import zipfile
from pathlib import Path
from typing import Any, Iterable
from lxml import etree as LET

try:
    import yaml
except ImportError as exc:  # pragma: no cover - exercised by installer checks
    raise SystemExit("PyYAML is required. Run the publication-formatting installer.") from exc

try:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Inches, Pt, RGBColor
except ImportError as exc:  # pragma: no cover - exercised by installer checks
    raise SystemExit("python-docx is required. Run the publication-formatting installer.") from exc


SKILL_ROOT = Path(__file__).resolve().parent.parent
ASSETS = SKILL_ROOT / "assets"
REFERENCES = SKILL_ROOT / "references"
EDITOR_BRIDGE = Path(__file__).with_name("editor_bridge.ps1")
SCHEMA_VERSION = 1

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
CP_NS = "http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
DC_NS = "http://purl.org/dc/elements/1.1/"
DCTERMS_NS = "http://purl.org/dc/terms/"
PKG_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"

REFUSAL_PATTERNS = {
    "unverified_original": r"\[UNVERIFIED CITATION\s*[—-]\s*NO ORIGINAL\]",
    "unverified_crosscheck": r"\[UNVERIFIED CITATION\s*[—-]\s*AI HAS NOT CROSS-CHECKED\]",
    "missing_locator": r"\[UNVERIFIED CITATION\s*[—-]\s*NO QUOTE OR PAGE LOCATOR\]",
    "anchor_none": r"<!--anchor:none:",
    "unsupported_claim": r"\[HIGH-WARN-CLAIM-NOT-SUPPORTED\]",
    "negative_constraint": r"\[HIGH-WARN-NEGATIVE-CONSTRAINT-VIOLATION",
    "fabricated_reference": r"\[HIGH-WARN-FABRICATED-REFERENCE\]",
    "anchorless_claim": r"\[HIGH-WARN-CLAIM-AUDIT-ANCHORLESS",
    "uncited_constraint": r"\[HIGH-WARN-CONSTRAINT-VIOLATION-UNCITED",
}

PRIVATE_FRONTMATTER_KEYS = {
    "author",
    "authors",
    "affiliation",
    "affiliations",
    "email",
    "emails",
    "corresponding_author",
    "orcid",
    "acknowledgments",
    "acknowledgements",
}


class FormatterError(RuntimeError):
    """Expected, user-actionable formatter failure."""


def utc_now() -> str:
    return dt.datetime.now(dt.timezone.utc).replace(microsecond=0).isoformat().replace("+00:00", "Z")


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def json_dump(data: Any, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def yaml_load(path: Path) -> dict[str, Any]:
    try:
        data = yaml.safe_load(path.read_text(encoding="utf-8"))
    except (OSError, UnicodeError, yaml.YAMLError) as exc:
        raise FormatterError(f"Cannot read YAML {path}: {exc}") from exc
    if not isinstance(data, dict):
        raise FormatterError(f"YAML root must be a mapping: {path}")
    return data


def strict_keys(data: dict[str, Any], allowed: set[str], label: str) -> None:
    unexpected = sorted(set(data) - allowed)
    if unexpected:
        raise FormatterError(f"Unknown {label} field(s): {', '.join(unexpected)}")


def resolve_inside(base: Path, value: str | None, label: str, *, required: bool = False) -> Path | None:
    if not value:
        if required:
            raise FormatterError(f"Missing required path: {label}")
        return None
    candidate = (base / value).resolve() if not Path(value).is_absolute() else Path(value).resolve()
    try:
        candidate.relative_to(base.resolve())
    except ValueError as exc:
        raise FormatterError(f"{label} must remain inside the project directory: {candidate}") from exc
    if required and not candidate.is_file():
        raise FormatterError(f"{label} does not exist: {candidate}")
    return candidate


def validate_profile(profile: dict[str, Any]) -> None:
    allowed = {
        "schema_version", "id", "name", "status", "source", "page", "body", "styles",
        "captions", "tables", "figures", "header", "footer", "references", "variants",
    }
    strict_keys(profile, allowed, "profile")
    for key in ("schema_version", "id", "name", "page", "body", "styles"):
        if key not in profile:
            raise FormatterError(f"Profile missing required field: {key}")
    if profile["schema_version"] != SCHEMA_VERSION:
        raise FormatterError(f"Unsupported profile schema_version: {profile['schema_version']}")
    if not re.fullmatch(r"[a-z0-9][a-z0-9-]{1,62}", str(profile["id"])):
        raise FormatterError("Profile id must be lowercase hyphen-case")
    page = profile["page"]
    if not isinstance(page, dict):
        raise FormatterError("profile.page must be a mapping")
    strict_keys(page, {"size", "orientation", "margins_cm", "header_distance_cm", "footer_distance_cm", "line_numbers"}, "page")
    if page.get("size") not in {"letter", "a4"}:
        raise FormatterError("profile.page.size must be letter or a4")
    margins = page.get("margins_cm")
    if not isinstance(margins, dict) or set(margins) != {"top", "bottom", "left", "right"}:
        raise FormatterError("profile.page.margins_cm must declare top/bottom/left/right")
    if any(not isinstance(v, (int, float)) or v <= 0 for v in margins.values()):
        raise FormatterError("All margins must be positive numbers")
    if not isinstance(profile["styles"], dict) or "normal" not in profile["styles"]:
        raise FormatterError("profile.styles.normal is required")
    supported_alignments = {"left", "center", "right", "justify"}
    supported_spacing = {"single", "onehalf", "double"}
    body = profile["body"]
    if not isinstance(body, dict):
        raise FormatterError("profile.body must be a mapping")
    style_configs = [("body", body), *[(f"styles.{name}", value) for name, value in profile["styles"].items()]]
    for label, style_config in style_configs:
        if not isinstance(style_config, dict):
            raise FormatterError(f"profile.{label} must be a mapping")
        alignment = style_config.get("alignment")
        if alignment is not None and alignment not in supported_alignments:
            raise FormatterError(f"profile.{label}.alignment must be left, center, right, or justify")
        spacing = style_config.get("line_spacing")
        if spacing is not None and not isinstance(spacing, (int, float)) and spacing not in supported_spacing:
            raise FormatterError(f"profile.{label}.line_spacing must be single, onehalf, double, or a number")


def validate_submission(config: dict[str, Any]) -> None:
    allowed = {
        "schema_version", "manuscript", "bibliography", "csl", "journal_profile", "template",
        "template_contract", "template_mode", "editor", "variants", "outputs", "output_root", "open_after",
        "desktop_copy", "blind_terms", "title_page", "supplementary",
    }
    strict_keys(config, allowed, "submission")
    for key in ("schema_version", "manuscript", "journal_profile"):
        if key not in config:
            raise FormatterError(f"Submission contract missing required field: {key}")
    if config["schema_version"] != SCHEMA_VERSION:
        raise FormatterError(f"Unsupported submission schema_version: {config['schema_version']}")
    if config.get("editor", "auto") not in {"auto", "word", "wps", "none"}:
        raise FormatterError("editor must be auto, word, wps, or none")
    variants = config.get("variants", ["anonymized"])
    if not variants or any(value not in {"anonymized", "full"} for value in variants):
        raise FormatterError("variants must contain anonymized and/or full")
    outputs = config.get("outputs", ["docx", "pdf"])
    if not outputs or any(value not in {"docx", "pdf"} for value in outputs):
        raise FormatterError("outputs must contain docx and/or pdf")
    if bool(config.get("template")) != bool(config.get("template_contract")):
        raise FormatterError("template and template_contract must be declared together")
    template_mode = config.get("template_mode", "profile_overlay")
    if template_mode not in {"profile_overlay", "template_authoritative"}:
        raise FormatterError("template_mode must be profile_overlay or template_authoritative")
    if template_mode == "template_authoritative" and not config.get("template"):
        raise FormatterError("template_authoritative mode requires template and template_contract")
    if config.get("supplementary") is not None and not isinstance(config["supplementary"], list):
        raise FormatterError("supplementary must be a list of project-relative files")


def command_version(path: str, args: list[str] | None = None) -> str | None:
    try:
        proc = subprocess.run([path, *(args or ["--version"])], capture_output=True, text=True, timeout=10, check=False)
    except (OSError, subprocess.TimeoutExpired):
        return None
    text = (proc.stdout or proc.stderr).strip().splitlines()
    return text[0] if text else None


def common_command(name: str, candidates: Iterable[Path]) -> str | None:
    for candidate in candidates:
        if candidate.is_file():
            return str(candidate.resolve())
    found = shutil.which(name)
    if found:
        return str(Path(found).resolve())
    return None


def windows_com_server(progid: str) -> dict[str, Any]:
    result: dict[str, Any] = {"progid": progid, "registered": False, "server": None}
    if os.name != "nt":
        return result
    try:
        import winreg
        with winreg.OpenKey(winreg.HKEY_CLASSES_ROOT, fr"{progid}\CLSID") as key:
            clsid = winreg.QueryValueEx(key, None)[0]
        with winreg.OpenKey(winreg.HKEY_CLASSES_ROOT, fr"CLSID\{clsid}\LocalServer32") as key:
            server = winreg.QueryValueEx(key, None)[0]
        result.update({"registered": True, "clsid": clsid, "server": server})
    except OSError:
        pass
    return result


def executable_from_server(server: str | None) -> str | None:
    if not server:
        return None
    quoted = re.match(r'^"([^"]+\.exe)"', server, flags=re.I)
    if quoted:
        return quoted.group(1)
    unquoted = re.match(r"^(.+?\.exe)(?:\s|$)", server, flags=re.I)
    return unquoted.group(1) if unquoted else None


def detect_editors() -> dict[str, Any]:
    if os.name != "nt":
        apps = Path("/Applications")
        word = apps / "Microsoft Word.app"
        wps = apps / "wpsoffice.app"
        return {
            "word": {"available": word.exists(), "path": str(word) if word.exists() else None, "automation": "not_checked"},
            "wps": {"available": wps.exists(), "path": str(wps) if wps.exists() else None, "automation": "not_checked"},
        }
    word_com = windows_com_server("Word.Application")
    wps_com = windows_com_server("KWPS.Application")
    word_exe = executable_from_server(word_com.get("server"))
    wps_exe = executable_from_server(wps_com.get("server"))
    genuine_word = bool(word_exe and Path(word_exe).name.lower() == "winword.exe" and Path(word_exe).is_file())
    wps_available = bool(wps_exe and Path(wps_exe).name.lower() == "wps.exe" and Path(wps_exe).is_file())
    return {
        "word": {
            "available": genuine_word,
            "path": word_exe if genuine_word else None,
            "registered": word_com.get("registered", False),
            "server": word_com.get("server"),
            "status": "available" if genuine_word else "not_checked",
            "reason": None if genuine_word else ("Word.Application is served by WPS" if word_exe and Path(word_exe).name.lower() == "wps.exe" else "Microsoft Word is not installed"),
        },
        "wps": {
            "available": wps_available,
            "path": wps_exe if wps_available else None,
            "registered": wps_com.get("registered", False),
            "server": wps_com.get("server"),
            "status": "available" if wps_available else "not_checked",
        },
    }


def detect_environment(*, self_test: bool = False) -> dict[str, Any]:
    program_files = Path(os.environ.get("ProgramFiles", ""))
    local_app_data = Path(os.environ.get("LOCALAPPDATA", ""))
    winget_pandoc = sorted((local_app_data / "Microsoft" / "WinGet" / "Packages").glob("JohnMacFarlane.Pandoc_*/*/pandoc.exe"))
    libreoffice_candidates = [
        program_files / "LibreOffice" / "program" / "soffice.com",
        program_files / "LibreOffice" / "program" / "soffice.exe",
        Path("/Applications/LibreOffice.app/Contents/MacOS/soffice"),
        Path("/usr/bin/libreoffice"),
    ]
    pandoc = common_command("pandoc", [
        local_app_data / "Pandoc" / "pandoc.exe",
        local_app_data / "Programs" / "Pandoc" / "pandoc.exe",
        program_files / "Pandoc" / "pandoc.exe",
        *winget_pandoc,
        Path("/usr/local/bin/pandoc"),
        Path("/opt/homebrew/bin/pandoc"),
    ])
    soffice = common_command("soffice", libreoffice_candidates)
    pdftoppm = common_command("pdftoppm", [
        Path.home() / ".cache" / "codex-runtimes" / "codex-primary-runtime" / "dependencies" / "native" / "poppler" / "Library" / "bin" / "pdftoppm.exe",
        Path("/usr/bin/pdftoppm"),
        Path("/opt/homebrew/bin/pdftoppm"),
    ])
    result = {
        "checked_at": utc_now(),
        "platform": platform.platform(),
        "python": {"path": sys.executable, "version": platform.python_version()},
        "pandoc": {"available": bool(pandoc), "path": pandoc, "version": command_version(pandoc) if pandoc else None},
        "libreoffice": {"available": bool(soffice), "path": soffice, "version": command_version(soffice) if soffice else None},
        "pdftoppm": {"available": bool(pdftoppm), "path": pdftoppm},
        "editors": detect_editors(),
        "self_test": "not_requested",
    }
    if self_test and soffice and pdftoppm:
        with tempfile.TemporaryDirectory(prefix="pubfmt-doctor-") as tmp:
            sample = Path(tmp) / "doctor.docx"
            doc = Document()
            doc.add_heading("Publication formatter self-test", level=1)
            doc.add_paragraph("The DOCX render pipeline is operational.")
            doc.save(sample)
            try:
                render = render_docx(sample, Path(tmp) / "render", soffice=soffice, pdftoppm=pdftoppm, timeout=60)
                result["self_test"] = {"status": "passed", "pages": len(render["pages"])}
            except FormatterError as exc:
                result["self_test"] = {"status": "failed", "error": str(exc)}
    return result


def split_front_matter(text: str) -> tuple[dict[str, Any], str]:
    if not text.startswith("---\n") and not text.startswith("---\r\n"):
        return {}, text
    match = re.match(r"^---\s*\r?\n(.*?)\r?\n---\s*\r?\n", text, flags=re.S)
    if not match:
        raise FormatterError("Malformed YAML front matter")
    metadata = yaml.safe_load(match.group(1)) or {}
    if not isinstance(metadata, dict):
        raise FormatterError("Manuscript front matter must be a mapping")
    return metadata, text[match.end():]


def manuscript_for_variant(source: Path, destination: Path, *, anonymized: bool) -> dict[str, Any]:
    text = source.read_text(encoding="utf-8")
    metadata, body = split_front_matter(text)
    if anonymized:
        metadata = {key: value for key, value in metadata.items() if str(key).lower() not in PRIVATE_FRONTMATTER_KEYS}
    rendered = "---\n" + yaml.safe_dump(metadata, sort_keys=False, allow_unicode=True).strip() + "\n---\n\n" + body if metadata else body
    destination.write_text(rendered, encoding="utf-8", newline="\n")
    return metadata


def scan_refusal_markers(text: str) -> list[dict[str, Any]]:
    findings: list[dict[str, Any]] = []
    for name, pattern in REFUSAL_PATTERNS.items():
        for match in re.finditer(pattern, text, flags=re.I):
            line = text.count("\n", 0, match.start()) + 1
            findings.append({"kind": name, "line": line, "text": match.group(0)})
    for match in re.finditer(r"<!--ref:[^>]+-->", text):
        marker = match.group(0)
        if "severity=HIGH-BLOCK" in marker:
            findings.append({"kind": "terminal_high_block", "line": text.count("\n", 0, match.start()) + 1, "text": marker})
        if not re.search(r"\b(ok|LOW-WARN-acknowledged)\b", marker) and not re.search(r"\bstatus=(ok|LOW-WARN-acknowledged)\b", marker):
            findings.append({"kind": "unresolved_ref_marker", "line": text.count("\n", 0, match.start()) + 1, "text": marker})
    return findings


def strip_ars_comments(text: str) -> str:
    return re.sub(r"<!--(?:ref|anchor|block):.*?-->", "", text, flags=re.S)


def semantic_manifest(markdown_text: str) -> dict[str, Any]:
    _, body = split_front_matter(markdown_text)
    clean = strip_ars_comments(body)
    headings = [{"level": len(m.group(1)), "text": re.sub(r"\s+#+\s*$", "", m.group(2)).strip()} for m in re.finditer(r"^(#{1,6})\s+(.+)$", clean, flags=re.M)]
    citations: list[str] = []
    for group in re.findall(r"\[([^\]]*@[A-Za-z0-9_:.+/-]+[^\]]*)\]", clean):
        citations.extend(re.findall(r"@([A-Za-z0-9_:.+/-]+)", group))
    images = [match[1] for match in re.findall(r"!\[([^\]]*)\]\(([^)]+)\)", clean)]
    table_delimiters = len(re.findall(r"^\s*\|?(?:\s*:?-+:?\s*\|){1,}", clean, flags=re.M))
    plain = re.sub(r"```.*?```", " ", clean, flags=re.S)
    plain = re.sub(r"`[^`]+`", " ", plain)
    plain = re.sub(r"!\[[^\]]*\]\([^)]+\)", " ", plain)
    plain = re.sub(r"\[[^\]]+\]\([^)]+\)", " ", plain)
    plain = re.sub(r"[#>*_~|]", " ", plain)
    words = re.findall(r"[A-Za-z0-9][A-Za-z0-9'’-]*|[\u3400-\u9fff]", plain)
    return {
        "headings": headings,
        "heading_count": len(headings),
        "citation_keys": sorted(set(citations)),
        "citation_occurrences": len(citations),
        "images": images,
        "image_count": len(images),
        "table_count": table_delimiters,
        "word_units": len(words),
    }


def set_font(style: Any, config: dict[str, Any]) -> None:
    font = style.font
    family = config.get("font", "Times New Roman")
    east_asia = config.get("east_asia_font", family)
    font.name = family
    if font._element.rPr is None:
        font._element.get_or_add_rPr()
    rfonts = font._element.rPr.get_or_add_rFonts()
    for attr, value in (("ascii", family), ("hAnsi", family), ("cs", family), ("eastAsia", east_asia)):
        rfonts.set(qn(f"w:{attr}"), value)
    if config.get("size_pt"):
        font.size = Pt(float(config["size_pt"]))
    for field in ("bold", "italic"):
        if field in config:
            setattr(font, field, bool(config[field]))
    color = config.get("color")
    if color:
        font.color.rgb = RGBColor.from_string(str(color).replace("#", "").upper())


def set_paragraph_style(style: Any, config: dict[str, Any]) -> None:
    set_font(style, config)
    fmt = style.paragraph_format
    if "space_before_pt" in config:
        fmt.space_before = Pt(float(config["space_before_pt"]))
    if "space_after_pt" in config:
        fmt.space_after = Pt(float(config["space_after_pt"]))
    if "first_line_indent_cm" in config:
        fmt.first_line_indent = Cm(float(config["first_line_indent_cm"]))
    if "left_indent_cm" in config:
        fmt.left_indent = Cm(float(config["left_indent_cm"]))
    if "hanging_indent_cm" in config:
        fmt.first_line_indent = Cm(-float(config["hanging_indent_cm"]))
        fmt.left_indent = Cm(float(config["hanging_indent_cm"]))
    spacing = config.get("line_spacing")
    if spacing == "single":
        fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    elif spacing == "onehalf":
        fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    elif spacing == "double":
        fmt.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    elif isinstance(spacing, (int, float)):
        fmt.line_spacing = float(spacing)
    alignment = config.get("alignment")
    if alignment:
        fmt.alignment = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }[alignment]
    if "keep_with_next" in config:
        fmt.keep_with_next = bool(config["keep_with_next"])


def ensure_style(document: Document, name: str, config: dict[str, Any]) -> None:
    try:
        style = document.styles[name]
    except KeyError:
        style = document.styles.add_style(name, 1)
    set_paragraph_style(style, config)


def add_field(paragraph: Any, instruction: str, cached: str = "1") -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = cached
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()._r
    for node in (begin, instr, separate, text, end):
        run.append(node)


def set_line_numbering(section: Any, enabled: bool) -> None:
    sect_pr = section._sectPr
    for existing in sect_pr.findall(qn("w:lnNumType")):
        sect_pr.remove(existing)
    if enabled:
        line = OxmlElement("w:lnNumType")
        line.set(qn("w:countBy"), "1")
        line.set(qn("w:restart"), "continuous")
        sect_pr.append(line)


def suppress_line_numbers(paragraph: Any) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:suppressLineNumbers")) is None:
        p_pr.append(OxmlElement("w:suppressLineNumbers"))


def set_update_fields(document: Document) -> None:
    settings = document.settings._element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        existing = OxmlElement("w:updateFields")
        settings.append(existing)
    existing.set(qn("w:val"), "true")


def set_table_borders(table: Any, style: str) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for child in list(borders):
        borders.remove(child)
    edges = ["top", "left", "bottom", "right", "insideH", "insideV"]
    for edge in edges:
        element = OxmlElement(f"w:{edge}")
        if style == "full_grid":
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "4")
        elif style == "three_line" and edge in {"top", "bottom"}:
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "8")
        elif style == "three_line" and edge == "insideH":
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "4")
        else:
            element.set(qn("w:val"), "nil")
        borders.append(element)


def set_cell_width(cell: Any, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_dxa))


def apply_table_geometry(table: Any, usable_dxa: int, config: dict[str, Any]) -> None:
    if not table.rows:
        return
    columns = len(table.rows[0].cells)
    weights = []
    for index in range(columns):
        length = max((len(row.cells[index].text.strip()) for row in table.rows if len(row.cells) > index), default=1)
        weights.append(min(max(length, 4), 60) + 12)
    total = sum(weights) or columns
    widths = [int(usable_dxa * weight / total) for weight in weights]
    widths[-1] += usable_dxa - sum(widths)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(usable_dxa))
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row_index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        if row_index == 0 and config.get("repeat_header", True):
            header = OxmlElement("w:tblHeader")
            header.set(qn("w:val"), "true")
            tr_pr.append(header)
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_table_borders(table, config.get("border_style", "three_line"))


def apply_profile(docx_path: Path, profile: dict[str, Any]) -> None:
    document = Document(docx_path)
    page = profile["page"]
    body = profile["body"]
    for section in document.sections:
        if page["size"] == "a4":
            section.page_width, section.page_height = Cm(21), Cm(29.7)
        else:
            section.page_width, section.page_height = Inches(8.5), Inches(11)
        orientation = page.get("orientation", "portrait")
        section.orientation = WD_ORIENT.LANDSCAPE if orientation == "landscape" else WD_ORIENT.PORTRAIT
        if orientation == "landscape" and section.page_width < section.page_height:
            section.page_width, section.page_height = section.page_height, section.page_width
        margins = page["margins_cm"]
        section.top_margin = Cm(float(margins["top"]))
        section.bottom_margin = Cm(float(margins["bottom"]))
        section.left_margin = Cm(float(margins["left"]))
        section.right_margin = Cm(float(margins["right"]))
        section.header_distance = Cm(float(page.get("header_distance_cm", 1.0)))
        section.footer_distance = Cm(float(page.get("footer_distance_cm", 1.0)))
        set_line_numbering(section, bool(page.get("line_numbers", False)))

    normal_cfg = {**body, **profile["styles"]["normal"]}
    ensure_style(document, "Normal", normal_cfg)
    role_map = {"title": "Title", "heading_1": "Heading 1", "heading_2": "Heading 2", "heading_3": "Heading 3", "caption": "Caption", "bibliography": "Bibliography"}
    for role, style_name in role_map.items():
        config = profile["styles"].get(role)
        if config:
            ensure_style(document, style_name, {**body, **config})

    header_cfg = profile.get("header", {})
    footer_cfg = profile.get("footer", {})
    for section in document.sections:
        if header_cfg.get("text"):
            paragraph = section.header.paragraphs[0]
            paragraph.text = str(header_cfg["text"])
            suppress_line_numbers(paragraph)
            paragraph.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER, "right": WD_ALIGN_PARAGRAPH.RIGHT}.get(header_cfg.get("alignment"), WD_ALIGN_PARAGRAPH.RIGHT)
            for run in paragraph.runs:
                run.font.name = body.get("font", "Times New Roman")
                run.font.size = Pt(float(header_cfg.get("size_pt", 8)))
        if footer_cfg.get("page_number", True):
            paragraph = section.footer.paragraphs[0]
            paragraph.clear()
            suppress_line_numbers(paragraph)
            paragraph.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER, "right": WD_ALIGN_PARAGRAPH.RIGHT}.get(footer_cfg.get("alignment"), WD_ALIGN_PARAGRAPH.CENTER)
            add_field(paragraph, "PAGE")

    usable_dxa = int((document.sections[0].page_width - document.sections[0].left_margin - document.sections[0].right_margin) / 635)
    for table in document.tables:
        apply_table_geometry(table, usable_dxa, profile.get("tables", {}))

    max_width_cm = float(profile.get("figures", {}).get("max_width_cm", usable_dxa / 567))
    max_width_emu = int(Cm(max_width_cm))
    for shape in document.inline_shapes:
        if shape.width > max_width_emu:
            ratio = shape.height / shape.width
            shape.width = max_width_emu
            shape.height = int(max_width_emu * ratio)

    set_update_fields(document)
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""
    document.core_properties.comments = ""
    document.save(docx_path)
    scrub_package_metadata(docx_path)


def prepare_reference_doc(template: Path, destination: Path) -> Path:
    """Create a privacy-safe local reference copy without mutating the retained source."""
    shutil.copy2(template, destination)
    document = Document(destination)
    for section in document.sections:
        for story in (section.header, section.first_page_header, section.even_page_header, section.footer, section.first_page_footer, section.even_page_footer):
            for paragraph in story.paragraphs:
                paragraph.clear()
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""
    document.core_properties.comments = ""
    document.save(destination)
    scrub_package_metadata(destination)
    return destination


def apply_template_postprocess(docx_path: Path, profile: dict[str, Any]) -> None:
    """Preserve uploaded template geometry/styles while adding safe release fields."""
    document = Document(docx_path)
    header_cfg = profile.get("header", {})
    footer_cfg = profile.get("footer", {})
    normal_style = document.styles["Normal"]
    font_name = normal_style.font.name or profile.get("body", {}).get("font", "Times New Roman")
    for section in document.sections:
        if header_cfg.get("text"):
            paragraph = section.header.paragraphs[0]
            paragraph.text = str(header_cfg["text"])
            suppress_line_numbers(paragraph)
            paragraph.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER, "right": WD_ALIGN_PARAGRAPH.RIGHT}.get(header_cfg.get("alignment"), WD_ALIGN_PARAGRAPH.RIGHT)
            for run in paragraph.runs:
                run.font.name = font_name
                run.font.size = Pt(float(header_cfg.get("size_pt", 8)))
        if footer_cfg.get("page_number", True):
            paragraph = section.footer.paragraphs[0]
            paragraph.clear()
            suppress_line_numbers(paragraph)
            paragraph.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER, "right": WD_ALIGN_PARAGRAPH.RIGHT}.get(footer_cfg.get("alignment"), WD_ALIGN_PARAGRAPH.CENTER)
            add_field(paragraph, "PAGE")
    set_update_fields(document)
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""
    document.core_properties.comments = ""
    document.save(docx_path)
    scrub_package_metadata(docx_path)


def theme_font_signature(docx_path: Path) -> dict[str, dict[str, str | None]]:
    """Return semantic theme-font mappings, ignoring harmless package reserialization."""
    with zipfile.ZipFile(docx_path, "r") as archive:
        theme_path = "word/theme/theme1.xml"
        if theme_path not in archive.namelist():
            return {}
        root = LET.fromstring(archive.read(theme_path))
    namespace = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
    result: dict[str, dict[str, str | None]] = {}
    for family in ("majorFont", "minorFont"):
        node = root.find(f".//a:fontScheme/a:{family}", namespaces=namespace)
        if node is None:
            continue
        values: dict[str, str | None] = {}
        for child in node:
            key = LET.QName(child).localname
            if key == "font":
                key = f"script:{child.get('script')}"
            values[key] = child.get("typeface")
        result[family] = values
    return result


def paragraph_style_signature(style: Any) -> dict[str, Any]:
    fmt = style.paragraph_format
    font_theme = None
    rpr = style.element.rPr
    if rpr is not None and rpr.rFonts is not None:
        font_theme = rpr.rFonts.get(qn("w:asciiTheme")) or rpr.rFonts.get(qn("w:hAnsiTheme"))
    line_spacing = fmt.line_spacing
    if hasattr(line_spacing, "pt"):
        line_spacing = round(line_spacing.pt, 3)
    elif line_spacing is not None:
        line_spacing = float(line_spacing)
    return {
        "name": style.name,
        "font": style.font.name,
        "font_theme": font_theme,
        "size_pt": round(style.font.size.pt, 3) if style.font.size else None,
        "bold": style.font.bold,
        "italic": style.font.italic,
        "color": str(style.font.color.rgb) if style.font.color and style.font.color.rgb else None,
        "alignment": int(fmt.alignment) if fmt.alignment is not None else None,
        "space_before_pt": round(fmt.space_before.pt, 3) if fmt.space_before else None,
        "space_after_pt": round(fmt.space_after.pt, 3) if fmt.space_after else None,
        "left_indent_in": round(fmt.left_indent / 914400, 4) if fmt.left_indent else None,
        "first_line_indent_in": round(fmt.first_line_indent / 914400, 4) if fmt.first_line_indent else None,
        "line_spacing": line_spacing,
        "keep_with_next": fmt.keep_with_next,
    }


def template_fidelity_report(docx_path: Path, contract: dict[str, Any]) -> dict[str, Any]:
    """Fail closed when an authoritative uploaded template did not control the output."""
    document = Document(docx_path)
    checks: list[dict[str, Any]] = []
    expected_theme_fonts = contract.get("theme_fonts", {})
    actual_theme_fonts = theme_font_signature(docx_path)
    theme_unchanged = bool(expected_theme_fonts) and actual_theme_fonts == expected_theme_fonts
    if expected_theme_fonts:
        checks.append({
            "id": "template_theme_fonts",
            "status": "pass" if theme_unchanged else "fail",
            "expected": expected_theme_fonts,
            "actual": actual_theme_fonts,
        })
    expected_sections = contract.get("sections", [])
    if expected_sections:
        expected = expected_sections[0]
        actual_section = document.sections[0]
        actual = {
            "page_width_in": round(actual_section.page_width / 914400, 4),
            "page_height_in": round(actual_section.page_height / 914400, 4),
            "margins_in": {
                "top": round(actual_section.top_margin / 914400, 4),
                "bottom": round(actual_section.bottom_margin / 914400, 4),
                "left": round(actual_section.left_margin / 914400, 4),
                "right": round(actual_section.right_margin / 914400, 4),
            },
        }
        page_ok = actual["page_width_in"] == expected["page_width_in"] and actual["page_height_in"] == expected["page_height_in"]
        margin_ok = all(abs(actual["margins_in"][key] - expected["margins_in"][key]) <= 0.01 for key in expected["margins_in"])
        checks.append({"id": "template_page_geometry", "status": "pass" if page_ok and margin_ok else "fail", "expected": expected, "actual": actual})
    expected_styles = {item["name"]: item for item in contract.get("styles", [])}
    for name in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3", "Caption", "Bibliography"):
        expected = expected_styles.get(name)
        if not expected:
            continue
        try:
            actual_style = document.styles[name]
        except KeyError:
            checks.append({"id": f"template_style_{name}", "status": "fail", "reason": "missing"})
            continue
        actual = paragraph_style_signature(actual_style)
        comparable = {key: value for key, value in expected.items() if key in actual and value is not None}
        def equivalent(key: str, expected_value: Any) -> bool:
            actual_value = actual[key]
            if actual_value == expected_value:
                return True
            if expected_value is False and actual_value is None:
                return True
            if key == "font" and actual_value is None:
                return bool(
                    theme_unchanged
                    and expected.get("font_theme")
                    and actual.get("font_theme") == expected.get("font_theme")
                )
            return False

        ok = all(equivalent(key, value) for key, value in comparable.items())
        checks.append({"id": f"template_style_{name}", "status": "pass" if ok else "fail", "expected": comparable, "actual": actual})
    return {"passed": bool(checks) and all(item["status"] == "pass" for item in checks), "checks": checks}


def apply_title_page_overrides(docx_path: Path) -> None:
    """Keep identifying title pages separate from anonymous running heads/line numbers."""
    document = Document(docx_path)
    for section in document.sections:
        set_line_numbering(section, False)
        for paragraph in section.header.paragraphs:
            paragraph.clear()
            suppress_line_numbers(paragraph)
        for paragraph in section.footer.paragraphs:
            suppress_line_numbers(paragraph)
    document.save(docx_path)
    scrub_package_metadata(docx_path)


def scrub_package_metadata(docx_path: Path) -> None:
    temp_path = docx_path.with_suffix(".scrub.tmp")
    with zipfile.ZipFile(docx_path, "r") as source, zipfile.ZipFile(temp_path, "w", compression=zipfile.ZIP_DEFLATED) as target:
        for info in source.infolist():
            if info.filename == "docProps/custom.xml" or info.filename.startswith("word/comments"):
                continue
            data = source.read(info.filename)
            if info.filename.endswith(".xml") or info.filename.endswith(".rels"):
                try:
                    root = LET.fromstring(data)
                    for element in root.iter():
                        for attr in list(element.attrib):
                            if attr.startswith(f"{{{W_NS}}}rsid"):
                                del element.attrib[attr]
                    for element in list(root.iter()):
                        for child in list(element):
                            if child.tag in {
                                f"{{{W_NS}}}commentRangeStart",
                                f"{{{W_NS}}}commentRangeEnd",
                                f"{{{W_NS}}}commentReference",
                            }:
                                element.remove(child)
                    if info.filename == "docProps/core.xml":
                        for tag in (f"{{{DC_NS}}}creator", f"{{{CP_NS}}}lastModifiedBy"):
                            node = root.find(tag)
                            if node is not None:
                                node.text = ""
                    if info.filename == "[Content_Types].xml":
                        for child in list(root):
                            part_name = child.attrib.get("PartName", "")
                            if part_name == "/docProps/custom.xml" or part_name.startswith("/word/comments"):
                                root.remove(child)
                    if info.filename.endswith(".rels"):
                        for child in list(root):
                            target_name = child.attrib.get("Target", "")
                            rel_type = child.attrib.get("Type", "")
                            if target_name == "docProps/custom.xml" or target_name.startswith("comments") or rel_type.endswith(("/comments", "/commentsExtended")):
                                root.remove(child)
                    data = LET.tostring(root, encoding="utf-8", xml_declaration=True)
                except LET.XMLSyntaxError:
                    pass
            target.writestr(info, data)
    os.replace(temp_path, docx_path)


def package_scan(docx_path: Path) -> dict[str, Any]:
    with zipfile.ZipFile(docx_path, "r") as archive:
        names = archive.namelist()
        story_names = [name for name in names if name.startswith("word/") and name.endswith(".xml")]
        story = b"\n".join(archive.read(name) for name in story_names)
        return {
            "parts": len(names),
            "comments_present": any(name.startswith("word/comments") for name in names),
            "tracked_insertions": len(re.findall(br"<w:ins\b", story)),
            "tracked_deletions": len(re.findall(br"<w:del\b", story)),
            "custom_properties_present": "docProps/custom.xml" in names,
            "rsid_attributes": len(re.findall(br"\bw:rsid[A-Za-z]*=", story)),
            "fields": sorted(set(match.decode("utf-8", errors="ignore").strip() for match in re.findall(br"<w:instrText[^>]*>(.*?)</w:instrText>", story, flags=re.S))),
        }


def docx_manifest(docx_path: Path) -> dict[str, Any]:
    document = Document(docx_path)
    headings = []
    paragraphs = []
    direct_runs = 0
    direct_paragraphs = 0
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            paragraphs.append(text)
        if paragraph.style and paragraph.style.name.startswith("Heading") and text:
            match = re.search(r"(\d+)$", paragraph.style.name)
            headings.append({"level": int(match.group(1)) if match else None, "text": text})
        ppr = paragraph._p.pPr
        direct_p_tags = {child.tag for child in ppr} - {qn("w:pStyle"), qn("w:numPr"), qn("w:sectPr")} if ppr is not None else set()
        if direct_p_tags:
            direct_paragraphs += 1
        for run in paragraph.runs:
            rpr = run._r.rPr
            direct_r_tags = {child.tag for child in rpr} - {qn("w:rStyle")} if rpr is not None else set()
            if direct_r_tags:
                direct_runs += 1
    return {
        "paragraph_count": len(paragraphs),
        "headings": headings,
        "heading_count": len(headings),
        "table_count": len(document.tables),
        "table_shapes": [[len(table.rows), len(table.columns)] for table in document.tables],
        "image_count": len(document.inline_shapes),
        "section_count": len(document.sections),
        "direct_formatting": {"runs": direct_runs, "paragraphs": direct_paragraphs},
        "package": package_scan(docx_path),
    }


def bibliography_keys(path: Path) -> set[str]:
    if path.suffix.lower() == ".json":
        data = json.loads(path.read_text(encoding="utf-8"))
        if not isinstance(data, list):
            raise FormatterError("CSL JSON bibliography must be a list")
        return {str(item["id"]) for item in data if isinstance(item, dict) and item.get("id")}
    if path.suffix.lower() == ".bib":
        return set(re.findall(r"@[A-Za-z]+\s*\{\s*([^,\s]+)", path.read_text(encoding="utf-8", errors="replace")))
    raise FormatterError("Bibliography must be .json or .bib")


def run_pandoc(source: Path, output: Path, *, bibliography: Path | None, csl: Path | None, reference_doc: Path, resource_path: Path) -> list[str]:
    env = detect_environment()
    pandoc = env["pandoc"]["path"]
    if not pandoc:
        raise FormatterError("Pandoc is required for deterministic citation rendering; run the installer first")
    command = [
        pandoc,
        str(source),
        "--from=markdown+yaml_metadata_block+citations+pipe_tables",
        "--to=docx",
        f"--output={output}",
        f"--reference-doc={reference_doc}",
        f"--resource-path={resource_path}",
    ]
    if bibliography:
        command.extend(["--citeproc", f"--bibliography={bibliography}"])
    if csl:
        command.append(f"--csl={csl}")
    try:
        proc = subprocess.run(command, capture_output=True, text=True, timeout=180, check=False)
    except subprocess.TimeoutExpired as exc:
        detail = (exc.stderr or exc.stdout or "").strip() if isinstance(exc.stderr or exc.stdout, str) else ""
        return {
            "status": "failed",
            "editor": editor,
            "error": f"Editor automation timed out after 180s" + (f": {detail}" if detail else ""),
            "command": command,
        }
    if proc.returncode != 0 or not output.is_file():
        raise FormatterError(f"Pandoc failed ({proc.returncode}): {(proc.stderr or proc.stdout).strip()}")
    return command


def build_reference_doc(profile: dict[str, Any], destination: Path) -> None:
    document = Document()
    document.add_paragraph("Reference title", style="Title")
    document.add_paragraph("Reference heading", style="Heading 1")
    document.add_paragraph("Reference body", style="Normal")
    if "Bibliography" not in [style.name for style in document.styles]:
        document.styles.add_style("Bibliography", 1)
    document.save(destination)
    apply_profile(destination, profile)


def proper_file_uri(path: Path) -> str:
    return path.resolve().as_uri()


def render_pdf_to_png(pdf: Path, output_dir: Path, pdftoppm: str | None = None) -> list[str]:
    pdftoppm = pdftoppm or detect_environment()["pdftoppm"]["path"]
    if not pdftoppm:
        raise FormatterError("pdftoppm is not available")
    output_dir.mkdir(parents=True, exist_ok=True)
    prefix = output_dir / "page"
    proc = subprocess.run([pdftoppm, "-png", "-r", "144", str(pdf), str(prefix)], capture_output=True, text=True, timeout=180, check=False)
    pages = sorted(output_dir.glob("page-*.png"))
    if proc.returncode != 0 or not pages:
        raise FormatterError(f"PDF rasterization failed: {(proc.stderr or proc.stdout).strip()}")
    return [str(path) for path in pages]


def render_docx(docx_path: Path, output_dir: Path, *, soffice: str | None = None, pdftoppm: str | None = None, timeout: int = 120) -> dict[str, Any]:
    environment = detect_environment()
    soffice = soffice or environment["libreoffice"]["path"]
    pdftoppm = pdftoppm or environment["pdftoppm"]["path"]
    if not soffice:
        raise FormatterError("LibreOffice/soffice is not available")
    if not pdftoppm:
        raise FormatterError("pdftoppm is not available")
    output_dir.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="pubfmt-lo-") as temp_name:
        temp = Path(temp_name)
        profile = temp / "profile"
        stage = temp / "stage"
        convert = temp / "convert"
        for directory in (profile, stage, convert):
            directory.mkdir(parents=True, exist_ok=True)
        staged_docx = stage / "input.docx"
        shutil.copy2(docx_path, staged_docx)
        command = [
            soffice,
            f"-env:UserInstallation={proper_file_uri(profile)}",
            "--headless",
            "--nologo",
            "--nodefault",
            "--nofirststartwizard",
            "--norestore",
            "--convert-to",
            "pdf",
            "--outdir",
            str(convert),
            str(staged_docx),
        ]
        env = os.environ.copy()
        env.update({"HOME": str(temp), "TMP": str(temp), "TEMP": str(temp), "SAL_USE_VCLPLUGIN": "svp"})
        try:
            proc = subprocess.run(command, capture_output=True, text=True, timeout=timeout, env=env, check=False)
        except subprocess.TimeoutExpired as exc:
            raise FormatterError(f"LibreOffice timed out after {timeout}s") from exc
        pdf = convert / "input.pdf"
        if proc.returncode != 0 or not pdf.is_file() or pdf.stat().st_size == 0:
            raise FormatterError(f"LibreOffice conversion failed ({proc.returncode}): {(proc.stderr or proc.stdout).strip()}")
        final_pdf = output_dir / f"{docx_path.stem}.pdf"
        shutil.copy2(pdf, final_pdf)
    page_dir = output_dir / "pages"
    pages = render_pdf_to_png(final_pdf, page_dir, pdftoppm)
    return {"status": "passed", "pdf": str(final_pdf), "pages": pages, "command": command}


def run_editor_bridge(input_docx: Path, output_docx: Path, output_pdf: Path, editor: str, *, visible: bool = False) -> dict[str, Any]:
    if os.name != "nt":
        return {"status": "not_checked", "editor": editor, "reason": "COM automation is Windows-only"}
    powershell = shutil.which("pwsh") or shutil.which("powershell")
    if not powershell or not EDITOR_BRIDGE.is_file():
        return {"status": "not_checked", "editor": editor, "reason": "PowerShell editor bridge is unavailable"}
    command = [
        powershell, "-NoProfile", "-ExecutionPolicy", "Bypass", "-File", str(EDITOR_BRIDGE),
        "-InputDocx", str(input_docx), "-OutputDocx", str(output_docx), "-OutputPdf", str(output_pdf),
        "-Editor", editor,
    ]
    if visible:
        command.append("-Visible")
    proc = subprocess.run(command, capture_output=True, text=True, timeout=180, check=False)
    lines = [line for line in proc.stdout.splitlines() if line.strip()]
    payload = None
    for line in reversed(lines):
        with contextlib.suppress(json.JSONDecodeError):
            payload = json.loads(line)
            break
    if proc.returncode != 0 or payload is None:
        return {"status": "failed", "editor": editor, "error": (proc.stderr or proc.stdout).strip(), "command": command}
    return payload


def verify_docx(docx_path: Path, *, source_manifest: dict[str, Any] | None = None, blind_terms: list[str] | None = None) -> dict[str, Any]:
    manifest = docx_manifest(docx_path)
    text = "\n".join(paragraph.text for paragraph in Document(docx_path).paragraphs)
    checks: list[dict[str, Any]] = []

    def add(check_id: str, passed: bool, detail: str) -> None:
        checks.append({"id": check_id, "status": "pass" if passed else "fail", "detail": detail})

    package = manifest["package"]
    add("no_comments", not package["comments_present"], f"comments_present={package['comments_present']}")
    add("no_tracked_changes", package["tracked_insertions"] == 0 and package["tracked_deletions"] == 0, f"insertions={package['tracked_insertions']}, deletions={package['tracked_deletions']}")
    add("metadata_scrubbed", not package["custom_properties_present"] and package["rsid_attributes"] == 0, f"custom_properties={package['custom_properties_present']}, rsid={package['rsid_attributes']}")
    unresolved = [token for token in ("TODO", "TBD", "[INSERT", "Error! Reference source not found") if token.lower() in text.lower()]
    add("no_placeholders", not unresolved, "none" if not unresolved else ", ".join(unresolved))
    blind_hits = [term for term in (blind_terms or []) if term and term.lower() in text.lower()]
    add("blind_terms_absent", not blind_hits, "none" if not blind_hits else ", ".join(blind_hits))
    if source_manifest:
        expected_headings = [(item["level"], item["text"]) for item in source_manifest.get("headings", [])]
        actual_headings = [(item["level"], item["text"]) for item in manifest["headings"]]
        add("headings_preserved", expected_headings == actual_headings, f"expected={len(expected_headings)}, actual={len(actual_headings)}")
        add("tables_preserved", source_manifest.get("table_count", 0) == manifest["table_count"], f"expected={source_manifest.get('table_count', 0)}, actual={manifest['table_count']}")
        add("images_preserved", source_manifest.get("image_count", 0) == manifest["image_count"], f"expected={source_manifest.get('image_count', 0)}, actual={manifest['image_count']}")
    return {"docx": str(docx_path), "sha256": sha256_file(docx_path), "manifest": manifest, "checks": checks, "passed": all(item["status"] == "pass" for item in checks)}


def make_run_id(profile_id: str, manuscript_hash: str) -> str:
    return f"{dt.datetime.now().strftime('%Y%m%dT%H%M%S')}-{profile_id}-{manuscript_hash[:8]}"


def copy_to_desktop(run_dir: Path) -> str | None:
    desktop = Path.home() / "Desktop"
    if not desktop.is_dir():
        return None
    destination = desktop / run_dir.name
    if destination.exists():
        destination = desktop / f"{run_dir.name}-{uuid.uuid4().hex[:6]}"
    shutil.copytree(run_dir, destination)
    return str(destination)


def load_template_contract(path: Path, template: Path) -> dict[str, Any]:
    contract = json.loads(path.read_text(encoding="utf-8"))
    expected = contract.get("reference", {}).get("sha256")
    actual = sha256_file(template)
    if not expected or expected != actual:
        raise FormatterError("Template hash does not match the distilled contract; distill it again")
    if contract.get("unresolved"):
        raise FormatterError("Template contract contains unresolved fidelity items")
    return contract


def format_submission(config_path: Path, *, editor_override: str | None = None, open_after: bool | None = None) -> dict[str, Any]:
    config_path = config_path.resolve()
    base = config_path.parent
    config = yaml_load(config_path)
    validate_submission(config)
    manuscript = resolve_inside(base, config["manuscript"], "manuscript", required=True)
    bibliography = resolve_inside(base, config.get("bibliography"), "bibliography", required=bool(config.get("bibliography")))
    csl = resolve_inside(base, config.get("csl"), "csl", required=bool(config.get("csl")))
    profile_path = resolve_inside(base, config["journal_profile"], "journal_profile", required=True)
    template = resolve_inside(base, config.get("template"), "template", required=bool(config.get("template")))
    contract_path = resolve_inside(base, config.get("template_contract"), "template_contract", required=bool(config.get("template_contract")))
    title_page = resolve_inside(base, config.get("title_page"), "title_page", required=bool(config.get("title_page")))
    supplementary = [resolve_inside(base, value, f"supplementary[{index}]", required=True) for index, value in enumerate(config.get("supplementary", []))]
    profile = yaml_load(profile_path)
    validate_profile(profile)
    template_mode = config.get("template_mode", "profile_overlay")
    template_contract = None
    if template and contract_path:
        template_contract = load_template_contract(contract_path, template)

    source_text = manuscript.read_text(encoding="utf-8")
    refusal = scan_refusal_markers(source_text)
    if refusal:
        raise FormatterError("Citation provenance hard gate refused formatting: " + "; ".join(f"{item['kind']}@L{item['line']}" for item in refusal))
    source_manifest = semantic_manifest(source_text)
    if bibliography:
        missing = sorted(set(source_manifest["citation_keys"]) - bibliography_keys(bibliography))
        if missing:
            raise FormatterError("Citation key(s) missing from bibliography: " + ", ".join(missing))
    elif source_manifest["citation_keys"]:
        raise FormatterError("Manuscript contains citation keys but no bibliography is declared")

    output_root_value = config.get("output_root", "submission-output")
    output_root = resolve_inside(base, output_root_value, "output_root")
    if output_root is None:
        raise FormatterError("output_root cannot be empty")
    output_root.mkdir(parents=True, exist_ok=True)
    run_id = make_run_id(profile["id"], sha256_file(manuscript))
    run_dir = output_root / run_id
    if run_dir.exists():
        run_dir = output_root / f"{run_id}-{uuid.uuid4().hex[:6]}"
    run_dir.mkdir(parents=True)
    qa_dir = run_dir / "qa"
    qa_dir.mkdir()

    editor_requested = editor_override or config.get("editor", "auto")
    editor_state = detect_editors()
    if editor_requested == "auto":
        editor_selected = "word" if editor_state["word"]["available"] else ("wps" if editor_state["wps"]["available"] else "none")
    else:
        editor_selected = editor_requested
    if editor_selected in {"word", "wps"} and not editor_state[editor_selected]["available"]:
        raise FormatterError(f"Requested editor is unavailable: {editor_selected}")

    files: list[dict[str, Any]] = []
    variant_reports: dict[str, Any] = {}
    if supplementary:
        supplementary_dir = run_dir / "supplementary"
        supplementary_dir.mkdir()
        for source in supplementary:
            assert source is not None
            destination = supplementary_dir / source.name
            if destination.exists():
                destination = supplementary_dir / f"{source.stem}-{sha256_file(source)[:8]}{source.suffix}"
            shutil.copy2(source, destination)
            files.append({"role": "supplementary", "path": str(destination.relative_to(run_dir)), "sha256": sha256_file(destination)})
    for variant in config.get("variants", ["anonymized"]):
        variant_dir = run_dir / variant
        variant_dir.mkdir()
        with tempfile.TemporaryDirectory(prefix="pubfmt-build-") as temp_name:
            temp = Path(temp_name)
            clean_md = temp / "manuscript.md"
            manuscript_for_variant(manuscript, clean_md, anonymized=variant == "anonymized")
            clean_md.write_text(strip_ars_comments(clean_md.read_text(encoding="utf-8")), encoding="utf-8", newline="\n")
            reference_doc = None
            if template is not None:
                reference_doc = prepare_reference_doc(template, temp / "uploaded-reference.docx")
            else:
                reference_doc = temp / "reference.docx"
                build_reference_doc(profile, reference_doc)
            core_docx = variant_dir / f"manuscript-{variant}.docx"
            pandoc_command = run_pandoc(clean_md, core_docx, bibliography=bibliography, csl=csl, reference_doc=reference_doc, resource_path=base)
            if template_mode == "template_authoritative":
                apply_template_postprocess(core_docx, profile)
            else:
                apply_profile(core_docx, profile)
            structural = verify_docx(core_docx, source_manifest=source_manifest, blind_terms=config.get("blind_terms", []) if variant == "anonymized" else [])
            json_dump(structural, qa_dir / f"structural-{variant}.json")
            if not structural["passed"]:
                failed = ", ".join(item["id"] for item in structural["checks"] if item["status"] == "fail")
                raise FormatterError(f"Structural verification failed for {variant}: {failed}")
            template_fidelity = {"passed": True, "checks": [], "status": "not_applicable"}
            if template_mode == "template_authoritative" and template_contract:
                template_fidelity = template_fidelity_report(core_docx, template_contract)
                template_fidelity["status"] = "passed" if template_fidelity["passed"] else "failed"
                json_dump(template_fidelity, qa_dir / f"template-fidelity-{variant}.json")
                if not template_fidelity["passed"]:
                    failed = ", ".join(item["id"] for item in template_fidelity["checks"] if item["status"] == "fail")
                    raise FormatterError(f"Uploaded template fidelity failed for {variant}: {failed}")
            files.append({"role": f"core_{variant}", "path": str(core_docx.relative_to(run_dir)), "sha256": sha256_file(core_docx)})

            editor_report = {"status": "not_requested", "editor": editor_selected}
            reviewed_docx = None
            editor_pdf = None
            reviewed_structural = {"passed": True, "checks": [], "status": "not_applicable"}
            reviewed_template_fidelity = {"passed": True, "checks": [], "status": "not_applicable"}
            if editor_selected != "none":
                reviewed_docx = variant_dir / f"{editor_selected}-reviewed-{variant}.docx"
                editor_pdf = variant_dir / f"{editor_selected}-reviewed-{variant}.pdf"
                editor_report = run_editor_bridge(core_docx, reviewed_docx, editor_pdf, editor_selected, visible=bool(open_after if open_after is not None else config.get("open_after", False)))
                if editor_report.get("status") != "passed":
                    raise FormatterError(f"{editor_selected} automation failed: {editor_report.get('error') or editor_report.get('reason')}")
                scrub_package_metadata(reviewed_docx)
                reviewed_structural = verify_docx(reviewed_docx, source_manifest=source_manifest, blind_terms=config.get("blind_terms", []) if variant == "anonymized" else [])
                reviewed_structural["status"] = "passed" if reviewed_structural["passed"] else "failed"
                json_dump(reviewed_structural, qa_dir / f"structural-{editor_selected}-{variant}.json")
                if not reviewed_structural["passed"]:
                    failed = ", ".join(item["id"] for item in reviewed_structural["checks"] if item["status"] == "fail")
                    raise FormatterError(f"{editor_selected} reviewed copy failed structural verification: {failed}")
                if template_mode == "template_authoritative" and template_contract:
                    reviewed_template_fidelity = template_fidelity_report(reviewed_docx, template_contract)
                    reviewed_template_fidelity["status"] = "passed" if reviewed_template_fidelity["passed"] else "failed"
                    json_dump(reviewed_template_fidelity, qa_dir / f"template-fidelity-{editor_selected}-{variant}.json")
                    if not reviewed_template_fidelity["passed"]:
                        failed = ", ".join(item["id"] for item in reviewed_template_fidelity["checks"] if item["status"] == "fail")
                        raise FormatterError(f"{editor_selected} changed uploaded template geometry or core styles: {failed}")
                files.append({"role": f"{editor_selected}_reviewed_{variant}", "path": str(reviewed_docx.relative_to(run_dir)), "sha256": sha256_file(reviewed_docx)})
                if editor_pdf.is_file():
                    files.append({"role": f"{editor_selected}_pdf_{variant}", "path": str(editor_pdf.relative_to(run_dir)), "sha256": sha256_file(editor_pdf)})

            lo_report: dict[str, Any]
            try:
                lo_report = render_docx(reviewed_docx or core_docx, qa_dir / f"libreoffice-{variant}")
            except FormatterError as exc:
                lo_report = {"status": "failed", "error": str(exc)}
            json_dump(lo_report, qa_dir / f"render-libreoffice-{variant}.json")
            if lo_report.get("status") != "passed":
                raise FormatterError(f"LibreOffice visual rendering failed for {variant}: {lo_report.get('error')}")
            if "pdf" in config.get("outputs", ["docx", "pdf"]) and not editor_pdf:
                core_pdf = variant_dir / f"manuscript-{variant}.pdf"
                shutil.copy2(Path(lo_report["pdf"]), core_pdf)
                files.append({"role": f"core_pdf_{variant}", "path": str(core_pdf.relative_to(run_dir)), "sha256": sha256_file(core_pdf)})

            editor_render = {"status": "not_applicable"}
            if editor_pdf and editor_pdf.is_file():
                pages = render_pdf_to_png(editor_pdf, qa_dir / f"{editor_selected}-pdf-{variant}")
                editor_render = {"status": "passed", "pdf": str(editor_pdf), "pages": pages}
                json_dump(editor_render, qa_dir / f"render-{editor_selected}-{variant}.json")
            variant_reports[variant] = {
                "pandoc_command": pandoc_command,
                "editor": editor_report,
                "libreoffice_render": lo_report,
                "editor_render": editor_render,
                "structural": structural,
                "reviewed_structural": reviewed_structural,
                "template_fidelity": template_fidelity,
                "reviewed_template_fidelity": reviewed_template_fidelity,
            }

    if title_page:
        title_dir = run_dir / "title-page"
        title_dir.mkdir()
        with tempfile.TemporaryDirectory(prefix="pubfmt-title-") as temp_name:
            temp = Path(temp_name)
            clean_title = temp / "title-page.md"
            title_text = title_page.read_text(encoding="utf-8")
            title_refusal = scan_refusal_markers(title_text)
            if title_refusal:
                raise FormatterError("Title-page evidence gate refused formatting")
            clean_title.write_text(strip_ars_comments(title_text), encoding="utf-8", newline="\n")
            reference_doc = None
            if template is not None:
                reference_doc = prepare_reference_doc(template, temp / "uploaded-reference.docx")
            else:
                reference_doc = temp / "reference.docx"
                build_reference_doc(profile, reference_doc)
            core_title = title_dir / "title-page.docx"
            pandoc_command = run_pandoc(clean_title, core_title, bibliography=None, csl=None, reference_doc=reference_doc, resource_path=base)
            if template_mode == "template_authoritative":
                apply_template_postprocess(core_title, profile)
            else:
                apply_profile(core_title, profile)
            apply_title_page_overrides(core_title)
            title_structural = verify_docx(core_title)
            json_dump(title_structural, qa_dir / "structural-title-page.json")
            if not title_structural["passed"]:
                failed = ", ".join(item["id"] for item in title_structural["checks"] if item["status"] == "fail")
                raise FormatterError(f"Title page failed structural verification: {failed}")
            files.append({"role": "core_title_page", "path": str(core_title.relative_to(run_dir)), "sha256": sha256_file(core_title)})
            title_editor_report = {"status": "not_requested", "editor": editor_selected}
            title_editor_pdf = None
            title_reviewed = None
            if editor_selected != "none":
                title_reviewed = title_dir / f"{editor_selected}-reviewed-title-page.docx"
                title_editor_pdf = title_dir / f"{editor_selected}-reviewed-title-page.pdf"
                title_editor_report = run_editor_bridge(core_title, title_reviewed, title_editor_pdf, editor_selected, visible=False)
                if title_editor_report.get("status") != "passed":
                    raise FormatterError(f"{editor_selected} title-page automation failed: {title_editor_report.get('error')}")
                scrub_package_metadata(title_reviewed)
                files.append({"role": f"{editor_selected}_reviewed_title_page", "path": str(title_reviewed.relative_to(run_dir)), "sha256": sha256_file(title_reviewed)})
                files.append({"role": f"{editor_selected}_pdf_title_page", "path": str(title_editor_pdf.relative_to(run_dir)), "sha256": sha256_file(title_editor_pdf)})
            title_lo = render_docx(title_reviewed or core_title, qa_dir / "libreoffice-title-page")
            json_dump(title_lo, qa_dir / "render-libreoffice-title-page.json")
            if "pdf" in config.get("outputs", ["docx", "pdf"]) and not title_editor_pdf:
                title_pdf = title_dir / "title-page.pdf"
                shutil.copy2(Path(title_lo["pdf"]), title_pdf)
                files.append({"role": "core_pdf_title_page", "path": str(title_pdf.relative_to(run_dir)), "sha256": sha256_file(title_pdf)})
            title_editor_render = {"status": "not_applicable"}
            if title_editor_pdf:
                title_pages = render_pdf_to_png(title_editor_pdf, qa_dir / f"{editor_selected}-pdf-title-page")
                title_editor_render = {"status": "passed", "pdf": str(title_editor_pdf), "pages": title_pages}
                json_dump(title_editor_render, qa_dir / f"render-{editor_selected}-title-page.json")
            variant_reports["title-page"] = {
                "pandoc_command": pandoc_command,
                "editor": title_editor_report,
                "libreoffice_render": title_lo,
                "editor_render": title_editor_render,
                "structural": title_structural,
            }

    manifest = {
        "schema_version": SCHEMA_VERSION,
        "run_id": run_dir.name,
        "created_at": utc_now(),
        "status": "qa_pending_visual_inspection",
        "inputs": {
            "submission": {"path": str(config_path), "sha256": sha256_file(config_path)},
            "manuscript": {"path": str(manuscript), "sha256": sha256_file(manuscript)},
            "bibliography": {"path": str(bibliography), "sha256": sha256_file(bibliography)} if bibliography else None,
            "csl": {"path": str(csl), "sha256": sha256_file(csl)} if csl else None,
            "profile": {"path": str(profile_path), "sha256": sha256_file(profile_path), "id": profile["id"], "status": profile.get("status")},
            "template": {"path": str(template), "sha256": sha256_file(template)} if template else None,
            "template_contract": {"path": str(contract_path), "sha256": sha256_file(contract_path), "mode": template_mode} if contract_path else None,
            "title_page": {"path": str(title_page), "sha256": sha256_file(title_page)} if title_page else None,
            "supplementary": [{"path": str(path), "sha256": sha256_file(path)} for path in supplementary if path],
        },
        "source_semantics": source_manifest,
        "environment": detect_environment(),
        "editor_selected": editor_selected,
        "variants": variant_reports,
        "files": files,
        "desktop_copy_requested": bool(config.get("desktop_copy", False)),
        "desktop_copy": None,
    }
    json_dump(manifest, run_dir / "run-manifest.json")
    return {"run_dir": str(run_dir), "manifest": str(run_dir / "run-manifest.json"), "status": manifest["status"]}


def distill_template(template: Path, output_dir: Path, *, source_url: str | None = None, render: bool = True) -> dict[str, Any]:
    template = template.resolve()
    if not template.is_file() or template.suffix.lower() not in {".docx", ".dotx"}:
        raise FormatterError("Template must be an existing DOCX or DOTX")
    output_dir = output_dir.resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    before = sha256_file(template)
    document = Document(template)
    sections = []
    for index, section in enumerate(document.sections, 1):
        sections.append({
            "index": index,
            "page_width_in": round(section.page_width / 914400, 4),
            "page_height_in": round(section.page_height / 914400, 4),
            "margins_in": {
                "top": round(section.top_margin / 914400, 4), "bottom": round(section.bottom_margin / 914400, 4),
                "left": round(section.left_margin / 914400, 4), "right": round(section.right_margin / 914400, 4),
            },
            "different_first_page": bool(section.different_first_page_header_footer),
        })
    styles = [paragraph_style_signature(style) for style in document.styles if style.type == 1]
    package_parts = []
    with zipfile.ZipFile(template, "r") as archive:
        for info in archive.infolist():
            package_parts.append({"path": info.filename, "size": info.file_size, "sha256": hashlib.sha256(archive.read(info.filename)).hexdigest()})
    contract = {
        "schema_version": SCHEMA_VERSION,
        "reference": {"filename": template.name, "sha256": before, "source_url": source_url, "distilled_at": utc_now()},
        "sections": sections,
        "styles": styles,
        "theme_fonts": theme_font_signature(template),
        "structure": docx_manifest(template),
        "package_parts": package_parts,
        "mode": "style_reference",
        "editable_slots": [],
        "preserve_only": [item["path"] for item in package_parts if item["path"].startswith(("customXml/", "word/theme/", "word/numbering"))],
        "unresolved": ["multiple_sections_require_explicit_slot_mapping"] if len(sections) > 1 else [],
        "render": {"status": "not_requested"},
    }
    if render:
        try:
            contract["render"] = render_docx(template, output_dir / "reference-render")
        except FormatterError as exc:
            contract["render"] = {"status": "failed", "error": str(exc)}
    if before != sha256_file(template):
        raise FormatterError("Template changed during distillation")
    json_dump(contract, output_dir / "template-contract.json")
    artifact = textwrap.dedent(f"""\
    # Template distillation contract

    - Reference: `{template.name}`
    - SHA-256: `{before}`
    - Source: {source_url or 'not declared'}
    - Sections: {len(sections)}
    - Paragraph styles: {len(styles)}
    - Package parts: {len(package_parts)}
    - Render status: {contract['render']['status']}

    Exact section/style/package evidence is stored in `template-contract.json`.
    The reference remains authoritative and must retain the recorded SHA-256.
    """)
    (output_dir / "artifact.md").write_text(artifact, encoding="utf-8", newline="\n")
    return contract


def finalize_visual_qa(manifest_path: Path, *, reviewer: str, confirm_every_page: bool) -> dict[str, Any]:
    """Record explicit page-by-page inspection and only then release the run."""
    manifest_path = manifest_path.resolve()
    if not confirm_every_page:
        raise FormatterError("Finalization requires --confirm-every-page after inspecting every rendered page")
    if not reviewer.strip():
        raise FormatterError("A reviewer identity or agent label is required")
    try:
        manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        raise FormatterError(f"Cannot read run manifest: {exc}") from exc
    run_dir = manifest_path.parent
    if manifest.get("status") not in {"qa_pending_visual_inspection", "qa_passed"}:
        raise FormatterError(f"Run is not eligible for visual finalization: {manifest.get('status')}")
    inspected: dict[str, dict[str, int]] = {}
    for variant, report in manifest.get("variants", {}).items():
        inspected[variant] = {}
        for renderer in ("libreoffice_render", "editor_render"):
            render = report.get(renderer, {})
            if render.get("status") not in {"passed", "not_applicable"}:
                raise FormatterError(f"{variant}/{renderer} is not render-ready")
            pages = render.get("pages", [])
            if render.get("status") == "passed" and not pages:
                raise FormatterError(f"{variant}/{renderer} has no rendered pages")
            missing = [page for page in pages if not Path(page).is_file()]
            if missing:
                raise FormatterError(f"{variant}/{renderer} has missing page image(s)")
            inspected[variant][renderer] = len(pages)
    manifest["status"] = "qa_passed"
    manifest["visual_inspection"] = {
        "reviewer": reviewer.strip(),
        "completed_at": utc_now(),
        "every_page_confirmed": True,
        "page_counts": inspected,
    }
    json_dump(manifest, manifest_path)
    if manifest.get("desktop_copy_requested") and not manifest.get("desktop_copy"):
        manifest["desktop_copy"] = copy_to_desktop(run_dir)
        json_dump(manifest, manifest_path)
    return {
        "manifest": str(manifest_path),
        "status": manifest["status"],
        "visual_inspection": manifest["visual_inspection"],
        "desktop_copy": manifest.get("desktop_copy"),
    }


def print_human_doctor(report: dict[str, Any]) -> None:
    print("Publication formatting doctor")
    print(f"- Python: {report['python']['version']} ({report['python']['path']})")
    for key in ("pandoc", "libreoffice", "pdftoppm"):
        item = report[key]
        print(f"- {key}: {'available' if item['available'] else 'missing'}" + (f" ({item.get('path')})" if item.get("path") else ""))
    for key in ("word", "wps"):
        item = report["editors"][key]
        suffix = f"; {item.get('reason')}" if item.get("reason") else ""
        print(f"- {key}: {'available' if item['available'] else 'not_checked'}{suffix}")
    if report["self_test"] != "not_requested":
        print(f"- render self-test: {report['self_test'].get('status')}")


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Deterministic Zotero/Obsidian/Codex publication formatter")
    sub = parser.add_subparsers(dest="command", required=True)
    doctor = sub.add_parser("doctor", help="Detect document and editor capabilities")
    doctor.add_argument("--json", action="store_true", dest="as_json")
    doctor.add_argument("--self-test", action="store_true")
    distill = sub.add_parser("distill", help="Distill a retained DOCX/DOTX template")
    distill.add_argument("--template", required=True, type=Path)
    distill.add_argument("--output-dir", required=True, type=Path)
    distill.add_argument("--source-url")
    distill.add_argument("--no-render", action="store_true")
    format_cmd = sub.add_parser("format", help="Create a new submission run")
    format_cmd.add_argument("--config", required=True, type=Path)
    format_cmd.add_argument("--editor", choices=["auto", "word", "wps", "none"])
    format_cmd.add_argument("--open", action="store_true", dest="open_after")
    verify = sub.add_parser("verify", help="Run structural verification on an existing DOCX")
    verify.add_argument("--docx", required=True, type=Path)
    verify.add_argument("--source-manifest", type=Path)
    verify.add_argument("--blind-term", action="append", default=[])
    verify.add_argument("--render-dir", type=Path)
    finalize = sub.add_parser("finalize", help="Release a run after explicit page-by-page visual inspection")
    finalize.add_argument("--manifest", required=True, type=Path)
    finalize.add_argument("--reviewer", required=True)
    finalize.add_argument("--confirm-every-page", action="store_true")
    return parser


def main(argv: list[str] | None = None) -> int:
    args = build_parser().parse_args(argv)
    try:
        if args.command == "doctor":
            report = detect_environment(self_test=args.self_test)
            if args.as_json:
                print(json.dumps(report, ensure_ascii=False, indent=2))
            else:
                print_human_doctor(report)
            return 0
        if args.command == "distill":
            report = distill_template(args.template, args.output_dir, source_url=args.source_url, render=not args.no_render)
            print(json.dumps({"contract": str(args.output_dir.resolve() / 'template-contract.json'), "render": report["render"]}, ensure_ascii=False))
            return 0 if report["render"].get("status") in {"passed", "not_requested"} else 2
        if args.command == "format":
            report = format_submission(args.config, editor_override=args.editor, open_after=args.open_after)
            print(json.dumps(report, ensure_ascii=False))
            return 0
        if args.command == "verify":
            source = json.loads(args.source_manifest.read_text(encoding="utf-8")) if args.source_manifest else None
            report = verify_docx(args.docx.resolve(), source_manifest=source, blind_terms=args.blind_term)
            if args.render_dir:
                report["render"] = render_docx(args.docx.resolve(), args.render_dir.resolve())
            print(json.dumps(report, ensure_ascii=False, indent=2))
            return 0 if report["passed"] else 2
        if args.command == "finalize":
            report = finalize_visual_qa(args.manifest, reviewer=args.reviewer, confirm_every_page=args.confirm_every_page)
            print(json.dumps(report, ensure_ascii=False, indent=2))
            return 0
    except (FormatterError, OSError, ValueError, json.JSONDecodeError) as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 2
    return 2


if __name__ == "__main__":
    raise SystemExit(main())
